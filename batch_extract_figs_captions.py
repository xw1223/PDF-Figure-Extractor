import os
import re
import sys
import csv
import html
import difflib
from pathlib import Path
from typing import List, Tuple, Optional

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# ========= USER CONFIG (can be overridden by command line args) =========
PDF_ROOT = r"/path/to/your/root_folder"          # The main folder containing many subfolders, each with one PDF
CITATIONS_TXT = r"/path/to/your/Citations.txt"   # A text file: one citation per line
OUTPUT_DOCX = r"/path/to/output/All_Figures_Captions.docx"
OUTPUT_CSV = r"/path/to/output/pdf_to_citation_map.csv"   # optional

# Minimum image size (pixels) to keep (filter out small logos/icons)
MIN_W = 500
MIN_H = 500
MIN_AREA = 300_000

# Fuzzy match threshold (0–1). Higher = stricter match.
MATCH_THRESHOLD = 0.55


# ========= HELPER FUNCTIONS =========
def strip_html_tags(s: str) -> str:
    """Remove basic HTML tags and decode entities like &amp;"""
    s = re.sub(r"<[^>]+>", "", s)
    return html.unescape(s).strip()


def normalize_for_match(s: str) -> str:
    """Normalize text for fuzzy comparison."""
    s = s.lower()
    s = re.sub(r"[_\-–—:,.;/\\()\[\]{}<>|!?\"'`~^*+=]+", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def read_citations(txt_path: str) -> List[str]:
    """Read all citation lines from the text file."""
    out = []
    with open(txt_path, "r", encoding="utf-8") as f:
        for line in f:
            s = strip_html_tags(line.strip())
            if s:
                out.append(s)
    return out


def best_fuzzy_match(target: str, pool: List[str], threshold: float) -> Tuple[str, float]:
    """Find the best fuzzy match of `target` in `pool`."""
    if not target:
        return "", 0.0
    target_n = normalize_for_match(target)
    best, best_score = "", 0.0
    for cand in pool:
        score = difflib.SequenceMatcher(None, target_n, normalize_for_match(cand)).ratio()
        if score > best_score:
            best, best_score = cand, score
    return (best, best_score) if best_score >= threshold else ("", 0.0)


def extract_pdf_title(doc: fitz.Document) -> Optional[str]:
    """Extract title from metadata or first-page largest text spans."""
    meta_title = (doc.metadata or {}).get("title") or ""
    if meta_title.strip():
        return meta_title.strip()

    try:
        page = doc[0]
        raw = page.get_text("dict")
        spans = []
        for block in raw.get("blocks", []):
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    txt = span.get("text", "").strip()
                    size = float(span.get("size", 0) or 0)
                    if txt:
                        spans.append((size, txt))
        if spans:
            spans.sort(key=lambda x: x[0], reverse=True)
            top_size = spans[0][0]
            pieces = [t for s, t in spans if abs(s - top_size) < 0.2]
            candidate = re.sub(r"\s+", " ", " ".join(pieces)).strip()
            bad_starts = (
                "Graphical abstract", "Highlights", "Article", "OPEN ACCESS",
                "Summary", "In brief", "STAR★METHODS", "REFERENCES"
            )
            if candidate and not any(candidate.startswith(b) for b in bad_starts):
                return candidate
    except Exception:
        pass
    return None


def extract_full_text(doc: fitz.Document) -> str:
    """Concatenate plain text from all pages."""
    return "\n".join(page.get_text("text") for page in doc)


def extract_captions(full_text: str) -> List[str]:
    """Extract captions starting with 'Figure' or 'Fig.' across pages."""
    normalized = re.sub(r"\bFig\.\s*", "Figure ", full_text)
    pattern = re.compile(
        r"(Figure\s+[S]?\d+[A-Za-z]?(?:\.[A-Za-z])?\s*[\.:].*?)"
        r"(?=(?:\nFigure\s+[S]?\d+)|\nSTAR★METHODS|\nREFERENCES|\nArticle|\Z)",
        flags=re.DOTALL,
    )
    out = []
    for m in pattern.finditer(normalized):
        cap = m.group(1).strip()
        cap = re.sub(r"[ \t]+\n", "\n", cap)
        cap = re.sub(r"\n{3,}", "\n\n", cap)
        out.append(cap)
    return out


def extract_images(doc: fitz.Document, out_dir: Path) -> List[Path]:
    """Extract all large images (skip small icons) and return file paths."""
    out_dir.mkdir(parents=True, exist_ok=True)
    paths = []
    for pi, page in enumerate(doc, start=1):
        for ii, img in enumerate(page.get_images(full=True), start=1):
            xref = img[0]
            try:
                pix = fitz.Pixmap(doc, xref)
                if pix.colorspace and pix.colorspace.n == 4:
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                w, h, area = pix.width, pix.height, pix.width * pix.height
                if w >= MIN_W and h >= MIN_H and area >= MIN_AREA:
                    fn = f"p{pi:02d}_img{ii:02d}.png"
                    p = out_dir / fn
                    pix.save(p.as_posix())
                    paths.append(p)
            except Exception:
                continue
    return paths


def write_citation(doc: Document, name: str, matched: bool):
    """Insert the citation name before each figure/caption pair."""
    run = doc.add_paragraph().add_run(name if name else "[No citation matched]")
    run.bold = True
    if not matched:
        doc.paragraphs[-1].add_run("  [UNMATCHED]").italic = True


# ========= MAIN PIPELINE =========
def main(pdf_root: str, citations_txt: str, output_docx: str, output_csv: Optional[str] = None):
    # Recursively find all PDFs under the root directory
    pdf_files = sorted(Path(pdf_root).rglob("*.pdf"))
    if not pdf_files:
        print("No PDFs found under:", pdf_root)
        return

    citations = read_citations(citations_txt)

    # Prepare Word document
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    doc.styles["Normal"].font.size = Pt(11)
    doc.add_heading("Combined Figures & Captions", level=1)
    doc.add_paragraph(
        "Automatically generated. Each figure image is followed by its caption, "
        "and each pair is preceded by its matched citation name."
    )

    # Optional mapping table
    mapping_rows = [("pdf_path", "pdf_title_detected", "matched_citation", "score")]
    images_root = Path(output_docx).parent / f"_images_{Path(output_docx).stem}"

    for idx, pdf_path in enumerate(pdf_files, start=1):
        pdf = fitz.open(pdf_path.as_posix())

        # --- Detect title and match to citation ---
        detected_title = extract_pdf_title(pdf) or ""
        best_cit, score = best_fuzzy_match(detected_title, citations, MATCH_THRESHOLD)
        matched = bool(best_cit)
        citation_name = best_cit if matched else pdf_path.stem
        mapping_rows.append((str(pdf_path), detected_title or "N/A", best_cit or "N/A", f"{score:.2f}"))

        # --- Extract captions and images ---
        full_text = extract_full_text(pdf)
        captions = extract_captions(full_text)
        img_dir = images_root / pdf_path.stem
        images = extract_images(pdf, img_dir)

        # --- Write into DOCX ---
        if idx > 1:
            doc.add_page_break()
        doc.add_heading(f"[{idx}] {pdf_path.name}", level=2)

        pairs = max(len(images), len(captions))
        if pairs == 0:
            write_citation(doc, citation_name, matched)
            doc.add_paragraph("[No large figures or captions detected.]")
        else:
            for i in range(pairs):
                write_citation(doc, citation_name, matched)

                # Add image
                if i < len(images):
                    try:
                        p = doc.add_picture(images[i].as_posix(), width=Inches(6.5))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception:
                        try:
                            p = doc.add_picture(images[i].as_posix(), width=Inches(5.5))
                            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        except Exception:
                            doc.add_paragraph(f"[Image failed: {images[i].name}]")
                else:
                    doc.add_paragraph("[Image not found for this caption]")

                # Add caption
                if i < len(captions):
                    para = doc.add_paragraph()
                    m = re.match(r"^(Figure\s+[S]?\d+[A-Za-z]?(?:\.[A-Za-z])?)\s*([\.:].*)$",
                                 captions[i], flags=re.DOTALL)
                    if m:
                        r1 = para.add_run(m.group(1)); r1.bold = True
                        para.add_run(m.group(2))
                    else:
                        para.add_run(captions[i])
                else:
                    doc.add_paragraph("[Caption not detected for the above image]")

        # Optional info
        info = f"Title detected: {detected_title or 'N/A'} | Citation: {(best_cit or 'N/A')} (score={score:.2f})"
        doc.add_paragraph(info).italic = True

        pdf.close()

    # --- Save Word document ---
    Path(output_docx).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)
    print("Saved DOCX:", output_docx)

    # --- Save mapping CSV ---
    if output_csv:
        Path(output_csv).parent.mkdir(parents=True, exist_ok=True)
        with open(output_csv, "w", newline="", encoding="utf-8") as f:
            w = csv.writer(f)
            w.writerows(mapping_rows)
        print("Saved mapping CSV:", output_csv)


if __name__ == "__main__":
    # Allow command-line usage:
    # python batch_extract_figs_captions.py <pdf_root> <citations.txt> <out.docx> [out.csv]
    if len(sys.argv) >= 4:
        PDF_ROOT = sys.argv[1]
        CITATIONS_TXT = sys.argv[2]
        OUTPUT_DOCX = sys.argv[3]
        OUTPUT_CSV = sys.argv[4] if len(sys.argv) >= 5 else None
    main(PDF_ROOT, CITATIONS_TXT, OUTPUT_DOCX, OUTPUT_CSV)
